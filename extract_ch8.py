import docx
import os

doc = docx.Document('docs/source/ABLRFD Users Manual.docx')
paras = doc.paragraphs

SECTIONS = [
    (10491, 10510, '8.1', 'Example Problems', '8.1-example-problems'),
    (10510, 10677, '8.2', 'Example 1 (Design)', '8.2-example-1-design'),
    (10677, 10858, '8.3', 'Example 2 (Design)', '8.3-example-2-design'),
    (10858, 11020, '8.4', 'Example 3', '8.4-example-3'),
    (11020, 11209, '8.5', 'Example 4', '8.5-example-4'),
    (11209, 11396, '8.6', 'Example 5', '8.6-example-5'),
    (11396, 11560, '8.7', 'Example 6 (Design)', '8.7-example-6-design'),
    (11560, 11703, '8.8', 'Example 7', '8.8-example-7'),
    (11703, 11891, '8.9', 'Example 8', '8.9-example-8'),
    (11891, 12043, '8.10', 'Example 9 (Design)', '8.10-example-9-design'),
    (12043, 12177, '8.11', 'Example 10', '8.11-example-10'),
    (12177, 12356, '8.12', 'Example 11', '8.12-example-11'),
    (12356, 12515, '8.13', 'Example 12 (Design)', '8.13-example-12-design'),
    (12515, 12694, '8.14', 'Example 13 (Design)', '8.14-example-13-design'),
    (12694, 12891, '8.15', 'Example 14', '8.15-example-14'),
]

# Maps para index -> doc.tables index
# 10496 is Table-style caption; 10503 is Normal-style continuation caption
TABLE_PARA_MAP = {10496: 147, 10503: 148}


def is_courier(para):
    total_chars = sum(len(r.text) for r in para.runs)
    if total_chars == 0:
        return False
    courier_chars = sum(len(r.text) for r in para.runs if r.font.name == 'Courier New')
    return courier_chars / total_chars > 0.5


def table_to_md(tbl):
    rows = []
    for idx, row in enumerate(tbl.rows):
        cells = []
        for c in row.cells:
            txt = c.text.replace('\n', ' ').strip()
            txt = txt.replace('|', '\\|')
            cells.append(txt)
        rows.append('| ' + ' | '.join(cells) + ' |')
        if idx == 0:
            sep = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
            rows.append(sep)
    return '\n'.join(rows)


def convert_section(start_para, end_para, sec_num, sec_title):
    lines = []
    lines.append(f'## {sec_num} &emsp; {sec_title}')
    lines.append('')

    i = start_para
    while i < end_para:
        para = paras[i]
        style = para.style.name
        text = para.text

        # TABLE_PARA_MAP check overrides style-based routing (handles both Table and Normal styles)
        if i in TABLE_PARA_MAP:
            tbl = doc.tables[TABLE_PARA_MAP[i]]
            tab_parts = text.split('\t', 1)
            if len(tab_parts) == 2:
                lines.append(f'\n**{tab_parts[0].strip()}** {tab_parts[1].strip()}')
            else:
                lines.append(f'\n**{text.strip()}**')
            lines.append('')
            lines.append(table_to_md(tbl))
            lines.append('')
            i += 1
            continue

        if style in ('Heading 1', 'Heading 2'):
            i += 1
            continue

        if style == 'Figure':
            if text.strip():
                tab_parts = text.split('\t', 1)
                if len(tab_parts) == 2:
                    lines.append(f'\n**{tab_parts[0].strip()}** {tab_parts[1].strip()}')
                else:
                    lines.append(f'\n**{text.strip()}**')
                lines.append('')
            i += 1
            continue

        if style in ('Normal', 'Body Text Indent', 'Table'):
            if not text.strip():
                if lines and lines[-1] != '':
                    lines.append('')
                i += 1
                continue

            if is_courier(para):
                courier_lines = [text]
                j = i + 1
                while j < end_para:
                    next_para = paras[j]
                    next_style = next_para.style.name
                    if next_style in ('Normal', 'Body Text Indent') and is_courier(next_para):
                        courier_lines.append(next_para.text)
                        j += 1
                    elif next_style in ('Normal', 'Body Text Indent') and not next_para.text.strip():
                        # Empty line — check if next non-empty is also courier
                        k = j + 1
                        while k < end_para and not paras[k].text.strip():
                            k += 1
                        if k < end_para and paras[k].style.name in ('Normal', 'Body Text Indent') and is_courier(paras[k]):
                            for m in range(j, k):
                                courier_lines.append(paras[m].text)
                            j = k
                        else:
                            break
                    else:
                        break

                if lines and lines[-1] != '':
                    lines.append('')
                lines.append('```')
                for cl in courier_lines:
                    lines.append(cl.rstrip())
                lines.append('```')
                lines.append('')
                i = j
                continue
            else:
                t = text.replace('\t', ' ').strip()
                if t:
                    lines.append(t)
                    lines.append('')
                i += 1
                continue

        # Any other style
        if text.strip():
            lines.append(text.strip())
            lines.append('')
        i += 1

    return '\n'.join(lines)


out_dir = 'docs/source/chapter-08'
os.makedirs(out_dir, exist_ok=True)

index_content = '''```{raw} latex
\\clearpage
```

# Chapter 8 — Example Problems

```{toctree}
:maxdepth: 1

8.1-example-problems
8.2-example-1-design
8.3-example-2-design
8.4-example-3
8.5-example-4
8.6-example-5
8.7-example-6-design
8.8-example-7
8.9-example-8
8.10-example-9-design
8.11-example-10
8.12-example-11
8.13-example-12-design
8.14-example-13-design
8.15-example-14
```
'''

with open(out_dir + '/index.md', 'w', encoding='utf-8') as f:
    f.write(index_content)
print('Written: index.md')

for start_p, end_p, sec_num, sec_title, filename in SECTIONS:
    content = convert_section(start_p, end_p, sec_num, sec_title)
    filepath = out_dir + '/' + filename + '.md'
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f'Written: {filename}.md ({len(content)} chars)')

print('Done!')
