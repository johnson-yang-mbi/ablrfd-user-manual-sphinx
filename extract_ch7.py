import docx
import os

doc = docx.Document('docs/source/ABLRFD Users Manual.docx')
paras = doc.paragraphs

SECTIONS = [
    (6040, 6088, '7.1', 'General Output Information', '7.1-general-output-information'),
    (6088, 6101, '7.2', 'Cover Page', '7.2-cover-page'),
    (6101, 6212, '7.3', 'Input Data', '7.3-input-data'),
    (6212, 7834, '7.4', 'Results', '7.4-results'),
    (7834, 8325, '7.5', 'Intermediate Results', '7.5-intermediate-results'),
    (8325, 10467, '7.6', 'Formatted Output Tables', '7.6-formatted-output-tables'),
    (10467, 10473, '7.7', 'Specification Check Warnings', '7.7-specification-check-warnings'),
    (10473, 10486, '7.8', 'Specification Check Failures', '7.8-specification-check-failures'),
]

H3_MAP = {
    6042: '7.1.1', 6044: '7.1.2', 6047: '7.1.3', 6049: '7.1.4', 6071: '7.1.5', 6082: '7.1.6',
    6103: '7.3.1', 6160: '7.3.2', 6176: '7.3.3',
    6214: '7.4.1', 6307: '7.4.2', 6338: '7.4.3', 6340: '7.4.4', 6497: '7.4.5',
    6882: '7.4.6', 6900: '7.4.7', 7069: '7.4.8', 7225: '7.4.9', 7441: '7.4.10',
    7508: '7.4.11', 7718: '7.4.12',
    7835: '7.5.1', 7951: '7.5.2', 8025: '7.5.3', 8096: '7.5.4', 8131: '7.5.5', 8241: '7.5.6',
    8327: '7.6.1', 8366: '7.6.2', 8589: '7.6.3', 8945: '7.6.4', 9018: '7.6.5',
    9037: '7.6.6', 9107: '7.6.7', 9304: '7.6.8', 9340: '7.6.9', 9400: '7.6.10',
    9463: '7.6.11', 9529: '7.6.12', 9574: '7.6.13', 9596: '7.6.14', 9697: '7.6.15',
    9738: '7.6.16', 9804: '7.6.17', 9987: '7.6.18', 10083: '7.6.19', 10116: '7.6.20',
    10236: '7.6.21', 10316: '7.6.22', 10388: '7.6.23',
}

H4_MAP = {
    6499: '7.4.5.1', 6537: '7.4.5.2', 6700: '7.4.5.3',
    6910: '7.4.7.1', 6983: '7.4.7.2',
    7071: '7.4.8.1', 7210: '7.4.8.2',
    9806: '7.6.17.1', 9921: '7.6.17.2',
}

TABLE_PARA_MAP = {6079: 123, 6085: 124}


def is_courier(para):
    total_chars = sum(len(r.text) for r in para.runs)
    if total_chars == 0:
        return False
    courier_chars = sum(len(r.text) for r in para.runs if r.font.name == 'Courier New')
    return courier_chars / total_chars > 0.5


def table_to_md(tbl):
    rows = []
    for idx, row in enumerate(tbl.rows):
        cells = []
        for c in row.cells:
            txt = c.text.replace('\n', ' ').strip()
            txt = txt.replace('|', '\\|')
            cells.append(txt)
        rows.append('| ' + ' | '.join(cells) + ' |')
        if idx == 0:
            sep = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
            rows.append(sep)
    return '\n'.join(rows)


def convert_section(start_para, end_para, sec_num, sec_title):
    lines = []
    lines.append(f'## {sec_num} &emsp; {sec_title}')
    lines.append('')

    i = start_para
    while i < end_para:
        para = paras[i]
        style = para.style.name
        text = para.text

        if style in ('Heading 1', 'Heading 2'):
            i += 1
            continue

        if style == 'Heading 3':
            num = H3_MAP.get(i, '')
            heading_text = text.strip()
            if num:
                lines.append(f'\n### {num} {heading_text}')
            else:
                lines.append(f'\n### {heading_text}')
            lines.append('')
            i += 1
            continue

        if style == 'Heading 4':
            num = H4_MAP.get(i, '')
            heading_text = text.strip()
            if num:
                lines.append(f'\n#### {num} {heading_text}')
            else:
                lines.append(f'\n#### {heading_text}')
            lines.append('')
            i += 1
            continue

        if style == 'Table':
            tbl_idx = TABLE_PARA_MAP.get(i)
            if tbl_idx is not None:
                tbl = doc.tables[tbl_idx]
                tab_parts = text.split('\t', 1)
                if len(tab_parts) == 2:
                    lines.append(f'\n**{tab_parts[0].strip()}** {tab_parts[1].strip()}')
                else:
                    lines.append(f'\n**{text.strip()}**')
                lines.append('')
                lines.append(table_to_md(tbl))
                lines.append('')
            i += 1
            continue

        if style == 'Figure':
            if text.strip():
                tab_parts = text.split('\t', 1)
                if len(tab_parts) == 2:
                    lines.append(f'\n**{tab_parts[0].strip()}** {tab_parts[1].strip()}')
                else:
                    lines.append(f'\n**{text.strip()}**')
                lines.append('')
            i += 1
            continue

        if style == '1':
            if text.strip():
                item_text = text.replace('\t', ' ').strip()
                lines.append(item_text)
            i += 1
            continue

        if style == 'List Paragraph':
            if text.strip():
                lines.append(f'- {text.strip()}')
            i += 1
            continue

        if style in ('Normal', 'Body Text Indent'):
            if not text.strip():
                # Empty paragraph - add blank line separator only if last line wasn't already blank
                if lines and lines[-1] != '':
                    lines.append('')
                i += 1
                continue

            if is_courier(para):
                # Group consecutive courier paragraphs into a code block
                courier_lines = [text]
                j = i + 1
                while j < end_para:
                    next_para = paras[j]
                    next_style = next_para.style.name
                    if next_style in ('Normal', 'Body Text Indent') and is_courier(next_para):
                        courier_lines.append(next_para.text)
                        j += 1
                    elif next_style in ('Normal', 'Body Text Indent') and not next_para.text.strip():
                        # Empty line - check if next non-empty is also courier
                        k = j + 1
                        while k < end_para and not paras[k].text.strip():
                            k += 1
                        if k < end_para and paras[k].style.name in ('Normal', 'Body Text Indent') and is_courier(paras[k]):
                            # Include empty lines in code block
                            for m in range(j, k):
                                courier_lines.append(paras[m].text)
                            j = k
                        else:
                            break
                    else:
                        break

                if lines and lines[-1] != '':
                    lines.append('')
                lines.append('```')
                for cl in courier_lines:
                    lines.append(cl.rstrip())
                lines.append('```')
                lines.append('')
                i = j
                continue
            else:
                t = text.replace('\t', ' ').strip()
                if t:
                    lines.append(t)
                    lines.append('')
                i += 1
                continue

        # Any other style
        if text.strip():
            lines.append(text.strip())
            lines.append('')
        i += 1

    return '\n'.join(lines)


out_dir = 'docs/source/chapter-07'
os.makedirs(out_dir, exist_ok=True)

# Get chapter intro text (paras 6036-6039)
intro_parts = []
for i in range(6036, 6040):
    t = paras[i].text.strip()
    if t:
        intro_parts.append(t)
intro_text = '\n\n'.join(intro_parts)

index_content = '''```{raw} latex
\\clearpage
```

# Chapter 7 — Output Description

''' + intro_text + '''

```{toctree}
:maxdepth: 1

7.1-general-output-information
7.2-cover-page
7.3-input-data
7.4-results
7.5-intermediate-results
7.6-formatted-output-tables
7.7-specification-check-warnings
7.8-specification-check-failures
```
'''

with open(out_dir + '/index.md', 'w', encoding='utf-8') as f:
    f.write(index_content)
print('Written: index.md')

for start_p, end_p, sec_num, sec_title, filename in SECTIONS:
    content = convert_section(start_p, end_p, sec_num, sec_title)
    filepath = out_dir + '/' + filename + '.md'
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f'Written: {filename}.md ({len(content)} chars)')

print('Done!')
